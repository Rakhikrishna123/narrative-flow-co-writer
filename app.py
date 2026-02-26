import streamlit as st
import datetime
import json
import os
import base64
import ollama
import io
from docx import Document

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="Narrative Flow Co-Writer", layout="wide")

# ---------------- FILE HISTORY ----------------
HISTORY_FILE = "chat_history.json"

def save_history_to_file(history):
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=4)

def load_history_from_file():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []

# ---------------- SESSION STATE ----------------
if "messages" not in st.session_state:
    st.session_state.messages = []

if "history" not in st.session_state:
    st.session_state.history = load_history_from_file()

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if "username" not in st.session_state:
    st.session_state.username = ""

if "story_text" not in st.session_state:
    st.session_state.story_text = ""

if "avatar" not in st.session_state:
    st.session_state.avatar = None

# ---------------- LOGIN PAGE ----------------
if not st.session_state.logged_in:

    st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(-45deg, #1e1e2f, #2c3e50, #000000, #3a3a52);
        background-size: 400% 400%;
        animation: gradient 10s ease infinite;
    }
    @keyframes gradient {
        0% {background-position: 0% 50%;}
        50% {background-position: 100% 50%;}
        100% {background-position: 0% 50%;}
    }
    .login-card {
        background: rgba(255,255,255,0.1);
        padding: 40px;
        border-radius: 20px;
        backdrop-filter: blur(10px);
        text-align: center;
        color: white;
        max-width: 400px;
        margin: auto;
        margin-top: 150px;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="login-card">', unsafe_allow_html=True)
    st.markdown("## 🌙 Narrative Flow Login")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username and password:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.rerun()
        else:
            st.error("Enter username and password")

    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ---------------- BACKGROUND ----------------
def set_bg(image_file):
    try:
        with open(image_file, "rb") as f:
            encoded = base64.b64encode(f.read()).decode()
        st.markdown(f"""
        <style>
        .stApp {{
            background-image: url("data:image/jpg;base64,{encoded}");
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
        }}
        </style>
        """, unsafe_allow_html=True)
    except:
        pass

# ---------------- SIDEBAR ----------------
st.sidebar.success(f"Logged in as {st.session_state.username}")

if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.messages = []
    st.session_state.story_text = ""
    st.session_state.avatar = None
    st.rerun()

# User avatar upload
st.sidebar.markdown("### 👤 Profile Picture")
uploaded_file = st.sidebar.file_uploader("Upload your avatar", type=["png","jpg","jpeg"])
if uploaded_file is not None:
    st.session_state.avatar = uploaded_file
if st.session_state.avatar:
    st.sidebar.image(st.session_state.avatar, width=120)

st.sidebar.title("⚙ Story Controls")
genre = st.sidebar.selectbox("Genre", ["Fantasy", "Sci-Fi", "Mystery", "Romance", "Horror"])
mode = st.sidebar.selectbox("Writing Mode", ["Rewrite", "Continue", "Summarize", "Expand"])
tone = st.sidebar.selectbox("Tone", ["Emotional", "Dark", "Humorous", "Inspirational"])

genre_backgrounds = {
    "Fantasy": "dark_fantasy.jpg",
    "Sci-Fi": "scifi.jpg",
    "Mystery": "mystery.jpg",
    "Romance": "romance.jpg",
    "Horror": "horror.jpg"
}
set_bg(genre_backgrounds[genre])

# ---------------- CHAT HISTORY ----------------
st.sidebar.title("📜 Chat History")
custom_title = st.sidebar.text_input("Enter Chat Title")
if st.sidebar.button("💾 Save Current Chat") and st.session_state.messages:
    title = custom_title if custom_title else "Untitled Story"
    chat_data = {
        "title": title,
        "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "messages": st.session_state.messages.copy()
    }
    st.session_state.history.append(chat_data)
    save_history_to_file(st.session_state.history)
    st.sidebar.success("Chat saved!")

if st.session_state.history:
    titles = [chat["title"] + " (" + chat["time"] + ")" for chat in st.session_state.history]
    selected = st.sidebar.selectbox("Load Previous Chat", ["Select"] + titles)
    if selected != "Select":
        index = titles.index(selected)
        col1, col2 = st.sidebar.columns(2)
        if col1.button("Load"):
            st.session_state.messages = st.session_state.history[index]["messages"]
            st.sidebar.success("Chat loaded!")
        if col2.button("Delete"):
            st.session_state.history.pop(index)
            save_history_to_file(st.session_state.history)
            st.session_state.messages = []
            st.sidebar.success("Chat deleted!")
            st.rerun()

if st.sidebar.button("🗑 Clear Current Chat"):
    st.session_state.messages = []
    st.session_state.story_text = ""

# ---------------- DOCX FUNCTION ----------------
def create_docx_buffer(text, genre, tone):
    doc = Document()
    doc.add_heading("Narrative Flow Story", level=1)
    doc.add_paragraph(f"Genre: {genre}")
    doc.add_paragraph(f"Tone: {tone}")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------- WHATSAPP STYLE CSS WITH CUSTOM FONT ----------------
st.markdown("""
<style>
.chat-container { 
    display:flex; 
    flex-direction:column; 
    gap:15px; 
    margin-top:20px; 
    font-family: 'Arial', sans-serif; 
    font-size: 16px; 
}

.chat-row { 
    display:flex; 
    align-items:flex-end; 
    gap:10px; 
}

.user-row { justify-content:flex-end; }
.ai-row { justify-content:flex-start; }

.avatar { 
    width:40px; 
    height:40px; 
    border-radius:50%; 
    object-fit:cover; 
    background:#ddd; 
    display:flex; 
    align-items:center; 
    justify-content:center; 
    font-size:18px; 
}

.chat-user { 
    background-color:#25D366; 
    color:#fff5e6;  /* user font color */
    padding:12px 16px; 
    border-radius:18px 18px 0px 18px; 
    max-width:60%; 
    word-wrap: break-word; 
    font-family: 'Verdana', sans-serif; /* user font style */
}

.chat-ai { 
    background-color:#ffffff; 
    color:#1a1a1a;  /* AI font color */
    padding:12px 16px; 
    border-radius:18px 18px 18px 0px; 
    max-width:60%; 
    word-wrap: break-word; 
    font-family: 'Georgia', serif; /* AI font style */
}
</style>
""", unsafe_allow_html=True)

# ---------------- HEADER ----------------
st.markdown("# 🌙 Narrative Flow Co-Writer")
st.markdown(f"**User:** {st.session_state.username}  \n{genre} • {mode} • {tone}")

# ---------------- DISPLAY CHAT ----------------
st.markdown('<div class="chat-container">', unsafe_allow_html=True)
for msg in st.session_state.messages:
    if msg["role"] == "user":
        if st.session_state.avatar:
            avatar_html = f"<img src='data:image/png;base64,{base64.b64encode(st.session_state.avatar.getvalue()).decode()}' class='avatar'>"
        else:
            avatar_html = "<div class='avatar'>👤</div>"
        st.markdown(f"<div class='chat-row user-row'>{avatar_html}<div class='chat-user'>{msg['content']}</div></div>", unsafe_allow_html=True)
    else:
        st.markdown(f"<div class='chat-row ai-row'><div class='avatar'>🤖</div><div class='chat-ai'>{msg['content']}</div></div>", unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ---------------- AI FUNCTION ----------------
def simple_ai_reply(user_input, genre, mode, tone):
    system_prompt = f"""
    You are a creative fiction writer.
    Write a {tone.lower()} {genre.lower()} story like a novel chapter.
    Follow mode: {mode}.
    Only write the story
    Continuous paragraph only.
    """
    response = ollama.chat(
        model="tinyllama",
        options={"temperature":0.6, "num_predict":600, "top_p":0.9},
        messages=[{"role":"system","content":system_prompt},{"role":"user","content":user_input}]
    )
    return response["message"]["content"]

# ---------------- CHAT INPUT ----------------
prompt = st.chat_input("Write your story...")
if prompt:
    st.session_state.messages.append({"role":"user","content":prompt})
    if st.session_state.avatar:
        avatar_html = f"<img src='data:image/png;base64,{base64.b64encode(st.session_state.avatar.getvalue()).decode()}' class='avatar'>"
    else:
        avatar_html = "<div class='avatar'>👤</div>"
    st.markdown(f"<div class='chat-row user-row'>{avatar_html}<div class='chat-user'>{prompt}</div></div>", unsafe_allow_html=True)

    thinking = st.empty()
    thinking.markdown("<div class='chat-ai'>🤖 Thinking...</div>", unsafe_allow_html=True)

    response = simple_ai_reply(prompt, genre, mode, tone)

    thinking.markdown(f"<div class='chat-row ai-row'><div class='avatar'>🤖</div><div class='chat-ai'>{response}</div></div>", unsafe_allow_html=True)

    st.session_state.messages.append({"role":"assistant","content":response})
    st.session_state.story_text = "\n\n".join([m["content"] for m in st.session_state.messages if m["role"]=="assistant"])

# ---------------- DOWNLOAD ----------------
if st.session_state.story_text:
    st.subheader("📄 Download Your Story")
    docx_buffer = create_docx_buffer(st.session_state.story_text, genre, tone)
    st.download_button(
        label="Download as DOCX",
        data=docx_buffer,
        file_name="Narrative_Flow_Story.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )